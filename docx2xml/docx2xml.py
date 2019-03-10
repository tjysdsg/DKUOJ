# parse command line arguments
import argparse
parser = argparse.ArgumentParser(prog='docx2xml',description='Convert docx to xml for importing problems on DKUOJ')
parser.add_argument('filename',type=str, action='store', help='docx file to convert')

full_file_name = vars(parser.parse_args())['filename']
file_name = full_file_name.split('.')[0]

# the directory which contains this file
import os
dir_path = os.path.dirname(os.path.realpath(__file__))

print("Using templates in", dir_path)

# ElementTree for parsing xml template
import xml.etree.ElementTree as ET
tree = ET.parse(os.path.join(dir_path,'oj-template.xml'))
root = tree.getroot()

# python-docx for parsing docx
import docx
doc = docx.Document(full_file_name)

# content of a problem
section_content = {"Title":"",
"Level":"",
"Description":"",
"Input":"",
"Output":"",
"Sample Input":"",
"Sample Output":"",
"Test Input":"",
"Test Output":"",
"Hint":"",
"Source/Category":""
}

# parse paragraph by paragraph, but some content might contains multiple paragraphs
for i in range(len(doc.paragraphs)):
    if doc.paragraphs[i].text in section_content.keys():
        t = doc.paragraphs[i].text
    else:
        # add new line after a paragraph
        section_content[t] += '\n' + doc.paragraphs[i].text

root[1][0].text = section_content['Title']
root[1][3].text = section_content['Description']
root[1][4].text = section_content['Input']
root[1][5].text = section_content['Output']
root[1][6].text = section_content['Sample Input']
root[1][7].text = section_content['Sample Output']
root[1][8].text = section_content['Test Input']
root[1][9].text = section_content['Test Output']
root[1][10].text = section_content['Hint']

tree.write(file_name + '.xml')
